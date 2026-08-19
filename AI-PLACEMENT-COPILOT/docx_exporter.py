import os
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocxResumeExporter:
    """
    Generates ATS-optimized Word documents (.docx) following single-column,
    machine-readable formatting standards.
    """

    def __init__(self):
        self.primary_color = RGBColor(30, 41, 59) # Slate dark
        self.body_color = RGBColor(51, 65, 85) # Charcoal text
        self.accent_color = RGBColor(15, 23, 42) # Midnight black

    def create_resume_document(
        self,
        candidate_name: str,
        contact_info: str,
        summary: str,
        skills: List[str],
        experience_items: List[Dict[str, Any]],
        education_items: List[str],
        output_filepath: str = "tailored_resume.docx"
    ) -> str:
        """
        Builds and saves an ATS-compliant .docx resume.
        """
        doc = Document()

        # Set 0.75-inch standard margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # 1. Header (Name & Contact)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(candidate_name)
        name_run.font.name = "Arial"
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        name_run.font.color.rgb = self.accent_color

        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(contact_info)
        contact_run.font.name = "Arial"
        contact_run.font.size = Pt(9.5)
        contact_run.font.color.rgb = self.body_color

        # 2. Professional Summary
        if summary:
            self._add_section_heading(doc, "PROFESSIONAL SUMMARY")
            p = doc.add_paragraph()
            p_run = p.add_run(summary)
            p_run.font.name = "Arial"
            p_run.font.size = Pt(10)
            p_run.font.color.rgb = self.body_color

        # 3. Technical Skills
        if skills:
            self._add_section_heading(doc, "TECHNICAL SKILLS")
            p = doc.add_paragraph()
            p_run = p.add_run(" • ".join(skills))
            p_run.font.name = "Arial"
            p_run.font.size = Pt(10)
            p_run.font.color.rgb = self.body_color

        # 4. Work Experience & Optimized STAR Bullets
        if experience_items:
            self._add_section_heading(doc, "WORK EXPERIENCE")
            for item in experience_items:
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"{item.get('role', 'Role')} | {item.get('company', 'Company')}")
                title_run.font.name = "Arial"
                title_run.font.size = Pt(11)
                title_run.font.bold = True
                title_run.font.color.rgb = self.accent_color

                for bullet in item.get("bullets", []):
                    bullet_para = doc.add_paragraph(style="List Bullet")
                    b_run = bullet_para.add_run(bullet)
                    b_run.font.name = "Arial"
                    b_run.font.size = Pt(10)
                    b_run.font.color.rgb = self.body_color

        # 5. Education
        if education_items:
            self._add_section_heading(doc, "EDUCATION")
            for edu in education_items:
                edu_para = doc.add_paragraph(style="List Bullet")
                e_run = edu_para.add_run(edu)
                e_run.font.name = "Arial"
                e_run.font.size = Pt(10)
                e_run.font.color.rgb = self.body_color

        doc.save(output_filepath)
        return os.path.abspath(output_filepath)

    def _add_section_heading(self, doc: Document, title: str):
        """Creates a standardized ATS section header with border divider."""
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(3)
        run = h.add_run(title)
        run.font.name = "Arial"
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = self.accent_color


# --- Verification ---
if __name__ == "__main__":
    exporter = DocxResumeExporter()
    test_path = exporter.create_resume_document(
        candidate_name="Alex Mercer",
        contact_info="San Francisco, CA • alex.mercer@email.com • linkedin.com/in/alexmercer • github.com/alexmercer",
        summary="Results-driven Software Engineer with expertise in building scalable backend services and AI systems.",
        skills=["Python", "FastAPI", "PyTorch", "Docker", "PostgreSQL", "REST APIs", "Git"],
        experience_items=[
            {
                "role": "AI Backend Engineer",
                "company": "Tech Innovations Inc.",
                "bullets": [
                    "Architected high-throughput document parsing microservice in Python, reducing ingestion latency by 85%.",
                    "Integrated dense vector embedding search for real-time ATS candidate ranking."
                ]
            }
        ],
        education_items=["B.S. in Computer Science - University of California, Berkeley (2020 - 2024)"],
        output_filepath="test_output.docx"
    )
    print(f"✅ Generated sample .docx resume at: {test_path}")
